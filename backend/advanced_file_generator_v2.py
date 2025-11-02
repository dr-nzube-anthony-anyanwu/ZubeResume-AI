"""
Advanced File Generator V2
HTML-based approach for proper markdown to PDF/DOCX conversion
"""

import os
import re
import logging
import tempfile
from typing import Dict, Optional
from datetime import datetime
from pathlib import Path
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import pdfkit
# import weasyprint  # Commented out due to Windows dependency issues
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from xhtml2pdf import pisa
import difflib
from fuzzywuzzy import fuzz

logger = logging.getLogger(__name__)

class AdvancedFileGeneratorV2:
    """Advanced file generator using HTML as intermediate format"""
    
    def __init__(self, output_dir: str = "outputs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        # PDF options for wkhtmltopdf - optimized for 1-2 page resume
        self.pdf_options = {
            'page-size': 'A4',
            'margin-top': '0.5in',
            'margin-right': '0.6in',
            'margin-bottom': '0.5in',
            'margin-left': '0.6in',
            'encoding': "UTF-8",
            'no-outline': None,
            'enable-local-file-access': None
        }
    
    def clean_markdown_text(self, text: str) -> str:
        """Clean and prepare markdown text for conversion"""
        # Remove the "Here is the final..." prefix if present
        if "Here is the final, document-ready content" in text:
            parts = text.split("Here is the final, document-ready content")
            if len(parts) > 1:
                text = parts[1].strip()
                # Also remove any trailing text after the content
                if "This content is now perfectly structured" in text:
                    text = text.split("This content is now perfectly structured")[0].strip()
                if "that will generate perfect files:" in text:
                    text = text.split("that will generate perfect files:")[1].strip()
        
        logger.info(f"🧹 Starting text cleaning - input length: {len(text)}")
        
        # Remove all emojis and special characters that cause PDF issues
        emoji_pattern = re.compile("["
                                  u"\U0001F600-\U0001F64F"  # emoticons
                                  u"\U0001F300-\U0001F5FF"  # symbols & pictographs
                                  u"\U0001F680-\U0001F6FF"  # transport & map symbols
                                  u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
                                  u"\U00002600-\U000027BF"  # miscellaneous symbols
                                  u"\U0001f900-\U0001f9ff"  # supplemental symbols
                                  u"\U0001fa70-\U0001faff"  # symbols and pictographs extended-a
                                  "📍📧📞🔥📜💼🚀⚡✅❌🔧🎯🎉💾📄"  # specific emojis
                                  "]+", flags=re.UNICODE)
        text = emoji_pattern.sub('', text)
        
        # Structure the header properly
        text = self._restructure_header(text)
        logger.info("✅ Header restructuring complete")
        
        # Remove duplicate content blocks AGGRESSIVELY
        text = self._remove_duplicate_content_aggressive(text)
        logger.info("✅ Duplicate content removal complete")
        
        # Fix common issues
        text = re.sub(r'\*\*([^*]+)\*\*', r'**\1**', text)  # Fix bold formatting
        text = re.sub(r'##\s+([^\n]+)', r'## \1', text)  # Fix headers
        text = re.sub(r'\*\s+([^\n]+)', r'* \1', text)  # Fix bullet points
        text = re.sub(r'\+\s+([^\n]+)', r'* \1', text)  # Convert + to * for bullets
        
        # Convert main section headers to proper markdown h2 format
        # This ensures they get the blue color and bold styling
        section_headers = [
            'PROFESSIONAL SUMMARY',
            'CORE TECHNICAL SKILLS', 
            'PROJECT EXPERIENCE',
            'EDUCATION',
            'CERTIFICATIONS',
            'ACHIEVEMENTS',
            'REFERENCES'
        ]
        
        for header in section_headers:
            # Convert standalone header lines to markdown h2
            text = re.sub(rf'^{re.escape(header)}\s*$', f'## {header}', text, flags=re.MULTILINE)
            # Also handle headers that might have extra spacing
            text = re.sub(rf'^{re.escape(header)}\s+', f'## {header}\n', text, flags=re.MULTILINE)
        
        logger.info("✅ Section headers converted to markdown h2 format")
        
        # Fix certification formatting - put each certification and institution on one line
        text = self._fix_certification_formatting(text)
        
        # Optimize for page length - compact formatting
        text = self._optimize_for_page_length(text)
        
        # Add proper line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize line breaks
        
        logger.info(f"🎯 Text cleaning complete - output length: {len(text)}")
        return text.strip()
    
    def _remove_duplicate_content_aggressive(self, text: str) -> str:
        """Aggressively remove duplicate content blocks and keep only bulleted versions"""
        lines = text.split('\n')
        processed_lines = []
        i = 0
        
        logger.info(f"Starting aggressive duplicate removal - input has {len(lines)} lines")
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Check if this is a project title line
            if any(project in line for project in ['NzubeGlaucoDetect', 'NzubeCare', 'Vivian AI', 'Multi-Agent Healthcare']):
                logger.info(f"🎯 Found project title: {line}")
                # Add the project title
                processed_lines.append(lines[i])
                i += 1
                
                # Skip any role/tech lines immediately after project title
                role_tech_lines = []
                while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('*') and not lines[i].strip().startswith('-'):
                    role_tech_line = lines[i].strip()
                    # Keep role and tech lines
                    if any(keyword in role_tech_line for keyword in ['Engineer', 'Tech:', 'Developer', 'Lead', 'AI', 'Python']):
                        logger.info(f"📋 Keeping role/tech line: {role_tech_line}")
                        role_tech_lines.append(lines[i])
                    i += 1
                
                processed_lines.extend(role_tech_lines)
                
                # Now we need to identify and remove duplicates
                # Strategy: collect ALL remaining content for this project, then deduplicate
                all_content_lines = []
                project_content_start = i
                
                # Collect ALL content until next project or section
                while i < len(lines):
                    current_line = lines[i].strip()
                    
                    # Stop conditions
                    if (current_line.startswith('##') or 
                        any(project in current_line for project in ['NzubeGlaucoDetect', 'NzubeCare', 'Vivian AI', 'Multi-Agent Healthcare']) or
                        current_line.upper().startswith('EDUCATION') or 
                        current_line.upper().startswith('CERTIFICATIONS') or
                        current_line.upper().startswith('ACHIEVEMENTS')):
                        break
                    
                    all_content_lines.append(lines[i])
                    i += 1
                
                logger.info(f"📝 Collected {len(all_content_lines)} lines of content for deduplication")
                
                # Now perform aggressive deduplication
                deduplicated_content = self._deduplicate_project_content(all_content_lines)
                processed_lines.extend(deduplicated_content)
                
                # Add spacing after project
                processed_lines.append("")
                continue
            else:
                processed_lines.append(lines[i])
                i += 1
        
        logger.info(f"✅ Duplicate removal complete - output has {len(processed_lines)} lines")
        return '\n'.join(processed_lines)
    
    def _deduplicate_project_content(self, content_lines):
        """Deduplicate project content and prefer bulleted version"""
        if not content_lines:
            return []
        
        # Group content into blocks
        content_blocks = []
        current_block = []
        
        for line in content_lines:
            if line.strip() == "":
                if current_block:
                    content_blocks.append(current_block)
                    current_block = []
            else:
                current_block.append(line)
        
        if current_block:
            content_blocks.append(current_block)
        
        logger.info(f"🔍 Found {len(content_blocks)} content blocks")
        
        # Identify duplicates using similarity
        unique_blocks = []
        for block in content_blocks:
            block_text = '\n'.join([l.strip() for l in block if l.strip()])
            
            # Skip empty blocks
            if not block_text:
                continue
            
            # Check if this block is similar to any existing unique block
            is_duplicate = False
            better_version_exists = False
            
            for existing_idx, existing_block in enumerate(unique_blocks):
                existing_text = '\n'.join([l.strip() for l in existing_block if l.strip()])
                
                # Use fuzzy matching to detect duplicates
                similarity = fuzz.ratio(block_text.lower(), existing_text.lower())
                
                if similarity > 80:  # High similarity threshold
                    logger.info(f"🔄 Found similar content (similarity: {similarity}%)")
                    is_duplicate = True
                    
                    # Prefer the bulleted version
                    block_has_bullets = any(l.strip().startswith('*') or l.strip().startswith('-') for l in block)
                    existing_has_bullets = any(l.strip().startswith('*') or l.strip().startswith('-') for l in existing_block)
                    
                    if block_has_bullets and not existing_has_bullets:
                        # Replace existing with bulleted version
                        logger.info("✅ Replacing non-bulleted version with bulleted version")
                        unique_blocks[existing_idx] = block
                        better_version_exists = True
                    elif existing_has_bullets and not block_has_bullets:
                        # Keep existing bulleted version
                        logger.info("✅ Keeping existing bulleted version")
                        better_version_exists = True
                    else:
                        # Both have bullets or both don't - keep first one
                        logger.info("⏭️ Keeping first version")
                        better_version_exists = True
                    break
            
            # If not a duplicate, add it
            if not is_duplicate:
                unique_blocks.append(block)
                logger.info("➕ Added unique content block")
        
        # Flatten unique blocks back to lines
        result = []
        for block in unique_blocks:
            result.extend(block)
            result.append("")  # Add spacing between blocks
        
        logger.info(f"✨ Deduplication result: {len(unique_blocks)} unique blocks")
        return result
    
    def _restructure_header(self, text: str) -> str:
        """Restructure the header section to match user's image format exactly"""
        lines = text.split('\n')
        restructured_lines = []
        header_processed = False
        
        i = 0
        while i < len(lines) and not header_processed:
            line = lines[i].strip()
            
            # Look for the name (first line usually)
            if i == 0 or (i < 5 and ('Dr.' in line or any(name_word in line for name_word in ['Anthony', 'Anyanwu']))):
                # Clean and extract name
                name = line.replace('**', '').strip()
                if 'Dr.' in name:
                    restructured_lines.append(f"# {name}")
                    i += 1
                    
                    # Collect job titles (next few lines before contact info)
                    job_titles = []
                    while i < len(lines) and i < 10:
                        next_line = lines[i].strip()
                        if not next_line:
                            i += 1
                            continue
                        
                        # Stop if we hit contact info
                        if any(contact_indicator in next_line.lower() for contact_indicator in ['@', '+', 'nigeria', 'abuja', 'linkedin', 'github', '.com']):
                            break
                        
                        # Stop if we hit a section header
                        if next_line.upper().startswith('PROFESSIONAL') or next_line.startswith('##'):
                            break
                        
                        # This should be a job title
                        clean_title = next_line.replace('**', '').replace('*', '').strip()
                        if clean_title and len(clean_title) > 5:  # Avoid single words
                            job_titles.append(clean_title)
                        i += 1
                    
                    # Add job titles as one line
                    if job_titles:
                        # Join with commas for better formatting
                        combined_titles = " | ".join(job_titles)
                        restructured_lines.append(f"## {combined_titles}")
                    
                    # Collect ALL contact information and put on ONE line
                    contact_items = []
                    while i < len(lines) and i < 20:
                        contact_line = lines[i].strip()
                        if not contact_line:
                            i += 1
                            continue
                        
                        # Stop if we hit a section
                        if contact_line.upper().startswith('PROFESSIONAL') or contact_line.startswith('##'):
                            break
                        
                        # Clean contact line and add to collection
                        clean_contact = contact_line.replace('**', '').replace('*', '').strip()
                        if clean_contact and any(contact_indicator in clean_contact.lower() for contact_indicator in ['@', '+', 'nigeria', 'abuja', 'linkedin', 'github', '.com']):
                            contact_items.append(clean_contact)
                        i += 1
                    
                    # Combine ALL contact info into ONE horizontal line
                    if contact_items:
                        all_contact = " | ".join(contact_items)
                        restructured_lines.append(all_contact)
                    
                    restructured_lines.append("")  # Add spacing
                    header_processed = True
                    continue
                else:
                    restructured_lines.append(line)
                    i += 1
            else:
                restructured_lines.append(line)
                i += 1
                if line.upper().startswith('PROFESSIONAL') or line.startswith('##'):
                    header_processed = True
        
        # Add remaining lines
        while i < len(lines):
            restructured_lines.append(lines[i])
            i += 1
        
        return '\n'.join(restructured_lines)
    
    def _fix_certification_formatting(self, text: str) -> str:
        """Fix certification formatting to keep each certification and institution on one line"""
        lines = text.split('\n')
        fixed_lines = []
        in_certifications = False
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Check if we're entering certifications section
            if 'CERTIFICATIONS' in line.upper() or '## CERTIFICATIONS' in line:
                in_certifications = True
                fixed_lines.append(lines[i])
                i += 1
                continue
            
            # Check if we're leaving certifications section
            if in_certifications and (line.startswith('## ') or line.upper() in ['EDUCATION', 'ACHIEVEMENTS', 'REFERENCES']):
                in_certifications = False
            
            if in_certifications and line:
                # Handle certification entries
                if line.startswith('- ') or line.startswith('* '):
                    cert_line = line
                    
                    # Check if next line is an institution (not a bullet point)
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if next_line and not next_line.startswith('-') and not next_line.startswith('*') and not next_line.startswith('##'):
                            # Combine certification with institution on same line
                            cert_line = f"{line} - {next_line}"
                            i += 1  # Skip the next line since we combined it
                    
                    fixed_lines.append(cert_line)
                else:
                    fixed_lines.append(lines[i])
            else:
                fixed_lines.append(lines[i])
            
            i += 1
        
        logger.info("✅ Certification formatting fixed")
        return '\n'.join(fixed_lines)
    
    def _optimize_for_page_length(self, text: str) -> str:
        """Optimize content for 1-2 page length by using more compact formatting"""
        lines = text.split('\n')
        optimized_lines = []
        
        for line in lines:
            # Reduce excessive spacing between sections
            if line.strip() == '' and len(optimized_lines) > 0 and optimized_lines[-1].strip() == '':
                continue  # Skip duplicate empty lines
            
            # Compact bullet points (remove excessive line spacing)
            if line.startswith('- ') or line.startswith('* '):
                # Ensure bullet points are concise
                line = line.strip()
                if len(line) > 150:  # If bullet point is too long, try to make it more concise
                    # Keep important technical details but make it more compact
                    pass
            
            optimized_lines.append(line)
        
        logger.info("✅ Content optimized for page length")
        return '\n'.join(optimized_lines)
        lines = text.split('\n')
        processed_lines = []
        in_header = False
        contact_items = []
        
        for i, line in enumerate(lines):
            # Detect if we're in header section (first few lines typically contain name and contact)
            if i < 10 and (line.strip().startswith('#') or 
                          any(contact_word in line.lower() for contact_word in ['email', 'phone', 'linkedin', 'location', 'website'])):
                in_header = True
            
            # Collect contact info lines for horizontal formatting
            if in_header and any(contact_word in line.lower() for contact_word in ['email', 'phone', 'linkedin', 'location', 'website']):
                # Clean the contact line
                clean_line = line.strip().replace('**', '').replace('*', '')
                if clean_line:
                    contact_items.append(clean_line)
            elif contact_items and (line.strip().startswith('#') or not line.strip()):
                # End of contact section, format it horizontally
                if contact_items:
                    contact_line = ' | '.join(contact_items)
                    processed_lines.append(contact_line)
                    contact_items = []
                in_header = False
                processed_lines.append(line)
            elif not contact_items:
                processed_lines.append(line)
        
        # Handle remaining contact items at end
        if contact_items:
            contact_line = ' | '.join(contact_items)
            processed_lines.append(contact_line)
        
        return '\n'.join(processed_lines).strip()
    
    def markdown_to_html(self, markdown_text: str) -> str:
        """Convert markdown to well-formatted HTML"""
        try:
            # Clean the markdown first
            clean_text = self.clean_markdown_text(markdown_text)
            
            # Convert to HTML
            html = markdown.markdown(clean_text, extensions=['extra', 'nl2br'])
            
            # Add CSS styling
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: 'Calibri', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                        line-height: 1.3;
                        color: #333;
                        max-width: 8.5in;
                        margin: 0 auto;
                        padding: 0.4in;
                        font-size: 10pt;
                    }}
                    
                    /* Header Section - Exact match to user's image */
                    .header {{
                        text-align: center;
                        margin-bottom: 25px;
                        padding-bottom: 15px;
                        border-bottom: 1px solid #ddd;
                    }}
                    
                    /* Name - Large, bold, centered - VERY PROMINENT */
                    h1:first-of-type {{
                        font-size: 28pt;
                        font-weight: bold;
                        text-align: center;
                        margin: 0 0 10px 0;
                        color: #000;
                        letter-spacing: 1px;
                    }}
                    
                    /* Professional titles - horizontal line under name */
                    h2:first-of-type {{
                        font-size: 12pt;
                        font-weight: normal;
                        text-align: center;
                        margin: 8px 0 12px 0;
                        color: #333;
                        border: none;
                        text-transform: none;
                        letter-spacing: normal;
                        line-height: 1.3;
                    }}
                    
                    /* Contact information - centered single line */
                    p:first-of-type {{
                        text-align: center;
                        font-size: 10pt;
                        color: #555;
                        margin: 10px 0 0 0;
                        line-height: 1.4;
                    }}
                    
                    /* Section Headers */
                    h2 {{
                        font-size: 13pt;
                        font-weight: bold;
                        color: #2E86C1;
                        margin-top: 16px;
                        margin-bottom: 6px;
                        border-bottom: 2px solid #3498db;
                        padding-bottom: 2px;
                        text-transform: uppercase;
                        letter-spacing: 0.5px;
                    }}
                    
                    /* Project Titles - Make them stand out more */
                    h3 {{
                        font-size: 13pt;
                        font-weight: bold;
                        color: #2c3e50;
                        margin-top: 15px;
                        margin-bottom: 6px;
                        border-bottom: 1px solid #bbb;
                        padding-bottom: 2px;
                    }}
                    
                    /* Strong project titles in paragraphs */
                    p strong {{
                        font-size: 13pt;
                        font-weight: bold;
                        color: #2c3e50;
                        display: block;
                        margin-bottom: 4px;
                    }}
                    
                    /* Job/Education entries */
                    .job-entry {{
                        margin-bottom: 15px;
                    }}
                    
                    .job-header {{
                        display: flex;
                        justify-content: space-between;
                        align-items: baseline;
                        margin-bottom: 4px;
                    }}
                    
                    .job-title {{
                        font-weight: bold;
                        font-size: 11pt;
                    }}
                    
                    .job-date {{
                        font-style: italic;
                        font-size: 10pt;
                        color: #666;
                    }}
                    
                    .company {{
                        font-style: italic;
                        color: #555;
                        margin-bottom: 6px;
                    }}
                    
                    /* Lists */
                    ul {{
                        margin: 6px 0;
                        padding-left: 18px;
                    }}
                    
                    li {{
                        margin-bottom: 3px;
                        line-height: 1.3;
                    }}
                    
                    /* Paragraphs */
                    p {{
                        margin: 6px 0;
                        text-align: justify;
                    }}
                    
                    /* Strong/Bold text */
                    strong, b {{
                        font-weight: bold;
                        color: #2c3e50;
                    }}
                    
                    /* Print optimization */
                    @media print {{
                        body {{ 
                            margin: 0;
                            padding: 0.5in;
                        }}
                        
                        /* Override for print - ensure header stays centered */
                        h1:first-of-type, h2:first-of-type, p:first-of-type {{
                            text-align: center;
                        }}
                    }}
                </style>
                </style>
            </head>
            <body>
                {html}
            </body>
            </html>
            """
            
            return styled_html
            
        except Exception as e:
            logger.error(f"Error converting markdown to HTML: {e}")
            return f"<html><body><p>Error processing content: {str(e)}</p></body></html>"
    
    def html_to_pdf(self, html_content: str, output_path: str) -> bool:
        """Convert HTML to PDF using xhtml2pdf (better CSS support)"""
        try:
            # Clean unicode characters that may cause encoding issues
            import re
            # Replace problematic unicode characters with their ASCII equivalents
            cleaned_html = html_content
            unicode_replacements = {
                '\U0001f4cd': '●',  # Replace 📍 with bullet
                '\U0001f4e7': 'Email:',  # Replace 📧 with Email:
                '\U0001f4de': 'Phone:',  # Replace 📞 with Phone:
                '\U0001f310': 'Website:',  # Replace 🌐 with Website:
                '\u2022': '•',  # Keep bullet points
                '\u2013': '-',  # En dash
                '\u2014': '-',  # Em dash
                '\u201c': '"',  # Left double quote
                '\u201d': '"',  # Right double quote
                '\u2019': "'",  # Right single quote
            }
            
            for unicode_char, replacement in unicode_replacements.items():
                cleaned_html = cleaned_html.replace(unicode_char, replacement)
            
            # Remove any remaining non-ASCII characters as fallback
            cleaned_html = re.sub(r'[^\x00-\x7F]+', '', cleaned_html)
            
            try:
                # Convert to PDF using xhtml2pdf
                with open(output_path, "w+b") as result_file:
                    pisa_status = pisa.CreatePDF(
                        cleaned_html.encode('utf-8'),
                        dest=result_file,
                        encoding='utf-8'
                    )
                
                return not pisa_status.err
                
            except Exception as e:
                logger.error(f"xhtml2pdf error: {e}")
                # Fallback to ReportLab if xhtml2pdf fails
                return self._html_to_pdf_reportlab_fallback(cleaned_html, output_path)
                    
        except Exception as e:
            logger.error(f"Error converting HTML to PDF: {e}")
            return False
    
    def _html_to_pdf_reportlab_fallback(self, html_content: str, output_path: str) -> bool:
        """Fallback PDF generation using ReportLab"""
        try:
            # Parse HTML content and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=letter, topMargin=0.5*inch)
            styles = getSampleStyleSheet()
            story = []
            
            # Create custom styles
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            
            # Large name style
            name_style = ParagraphStyle(
                'NameStyle',
                parent=styles['Title'],
                fontSize=24,
                spaceAfter=8,
                alignment=TA_CENTER,
                textColor='black'
            )
            
            # Job titles style
            job_style = ParagraphStyle(
                'JobStyle',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=10,
                alignment=TA_CENTER,
                textColor='#333'
            )
            
            # Contact info style
            contact_style = ParagraphStyle(
                'ContactStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=15,
                alignment=TA_CENTER,
                textColor='#555'
            )
            
            # Project title style
            project_style = ParagraphStyle(
                'ProjectStyle',
                parent=styles['Heading3'],
                fontSize=13,
                fontName='Helvetica-Bold',
                spaceAfter=6,
                spaceBefore=12,
                textColor='#2c3e50'
            )
            
            element_count = 0
            
            # Process content elements
            for element in soup.find('body').children if soup.find('body') else []:
                if hasattr(element, 'name'):
                    if element.name == 'h1':
                        # Large header for name
                        story.append(Paragraph(element.get_text(), name_style))
                        element_count += 1
                        
                    elif element.name == 'h2':
                        # Job titles or section headers
                        if element_count <= 2:  # If this is early in document (job titles)
                            story.append(Paragraph(element.get_text(), job_style))
                        else:  # Section headers
                            section_style = ParagraphStyle(
                                'SectionStyle',
                                parent=styles['Heading2'],
                                fontSize=13,
                                fontName='Helvetica-Bold',
                                spaceAfter=6,
                                spaceBefore=12,
                                textColor='#2E86C1'  # Blue color to match CSS
                            )
                            story.append(Paragraph(element.get_text(), section_style))
                        element_count += 1
                            
                    elif element.name == 'h3':
                        # Project titles - make them bold and prominent
                        story.append(Paragraph(element.get_text(), project_style))
                        element_count += 1
                        
                    elif element.name == 'p':
                        text = element.get_text()
                        if element_count <= 3 and any(contact in text.lower() for contact in ['@', '+', 'nigeria', 'abuja']):
                            # Contact information
                            story.append(Paragraph(text, contact_style))
                        else:
                            # Regular paragraph
                            # Check if it contains project titles that should be bold
                            if any(project in text for project in ['NzubeGlaucoDetect', 'NzubeCare', 'Vivian AI', 'Multi-Agent Healthcare']):
                                story.append(Paragraph(text, project_style))
                            else:
                                normal_style = ParagraphStyle(
                                    'NormalStyle',
                                    parent=styles['Normal'],
                                    fontSize=10,
                                    spaceAfter=6,
                                    alignment=TA_LEFT
                                )
                                story.append(Paragraph(text, normal_style))
                        element_count += 1
                        
                    elif element.name == 'ul':
                        for li in element.find_all('li'):
                            bullet_style = ParagraphStyle(
                                'BulletStyle',
                                parent=styles['Normal'],
                                fontSize=10,
                                leftIndent=20,
                                spaceAfter=3
                            )
                            text = "• " + li.get_text()
                            story.append(Paragraph(text, bullet_style))
                        story.append(Spacer(1, 0.1*inch))
                        element_count += 1
            
            # Build PDF
            doc.build(story)
            return True
            
        except Exception as e:
            logger.error(f"Error with ReportLab fallback: {e}")
            return False
    
    def html_to_docx(self, html_content: str, output_path: str) -> bool:
        """Convert HTML to DOCX with improved content handling"""
        try:
            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Create new document
            doc = Document()
            
            # Set up styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Find the body content
            body = soup.find('body')
            if not body:
                # If no body tag, use the entire content
                body = soup
            
            # Process all elements in the body
            for element in body.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'div']):
                try:
                    text_content = element.get_text().strip()
                    if not text_content:
                        continue
                    
                    if element.name == 'h1':
                        p = doc.add_heading(text_content, level=1)
                        for run in p.runs:
                            run.font.size = Pt(18)
                            run.font.bold = True
                            
                    elif element.name == 'h2':
                        p = doc.add_heading(text_content, level=2)
                        # Make section headers blue and bold
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(46, 134, 193)  # Blue color
                            run.font.bold = True
                            run.font.size = Pt(13)
                            
                    elif element.name == 'h3':
                        p = doc.add_heading(text_content, level=3)
                        for run in p.runs:
                            run.font.size = Pt(12)
                            run.font.bold = True
                            
                    elif element.name in ['p', 'div']:
                        # Handle paragraphs and divs
                        if text_content:
                            p = doc.add_paragraph()
                            self._add_formatted_text_improved(p, element)
                            
                    elif element.name == 'ul':
                        for li in element.find_all('li'):
                            li_text = li.get_text().strip()
                            if li_text:
                                p = doc.add_paragraph(li_text, style='List Bullet')
                                
                    elif element.name == 'ol':
                        for li in element.find_all('li'):
                            li_text = li.get_text().strip()
                            if li_text:
                                p = doc.add_paragraph(li_text, style='List Number')
                                
                except Exception as elem_error:
                    logger.warning(f"Error processing element {element.name}: {elem_error}")
                    # Add as plain text if processing fails
                    if element.get_text().strip():
                        doc.add_paragraph(element.get_text().strip())
            
            # If document is empty, add the plain text content
            if len(doc.paragraphs) == 0:
                plain_text = soup.get_text().strip()
                if plain_text:
                    for line in plain_text.split('\n'):
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)
            
            # Save document
            doc.save(output_path)
            
            # Verify the file was created and has content
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                logger.info(f"DOCX file successfully created: {output_path} ({os.path.getsize(output_path)} bytes)")
                return True
            else:
                logger.error(f"DOCX file creation failed or empty: {output_path}")
                return False
            
        except Exception as e:
            logger.error(f"Error converting HTML to DOCX: {e}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return False

    def _add_formatted_text_improved(self, paragraph, element):
        """Add formatted text to paragraph from HTML element with better handling"""
        try:
            # Handle text content with formatting
            for content in element.contents:
                if hasattr(content, 'name') and content.name:
                    if content.name in ['strong', 'b']:
                        run = paragraph.add_run(content.get_text())
                        run.bold = True
                    elif content.name in ['em', 'i']:
                        run = paragraph.add_run(content.get_text())
                        run.italic = True
                    elif content.name == 'br':
                        paragraph.add_run('\n')
                    else:
                        text = content.get_text()
                        if text.strip():
                            paragraph.add_run(text)
                else:
                    # Plain text content
                    text = str(content).strip()
                    if text:
                        paragraph.add_run(text)
                        
        except Exception as e:
            # Fallback: add plain text
            try:
                text = element.get_text().strip()
                if text:
                    paragraph.add_run(text)
            except:
                logger.warning(f"Could not add formatted text: {e}")
    
    def generate_pdf_advanced(self, resume_text: str, filename: str = None, style: str = 'modern') -> Dict[str, str]:
        """Generate PDF using advanced HTML conversion"""
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"resume_{timestamp}.pdf"
            
            if not filename.endswith('.pdf'):
                filename += '.pdf'
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Convert markdown to HTML
            html_content = self.markdown_to_html(resume_text)
            
            # Convert HTML to PDF
            success = self.html_to_pdf(html_content, file_path)
            
            if success and os.path.exists(file_path):
                return {
                    'file_path': file_path,
                    'filename': filename,
                    'file_size': os.path.getsize(file_path),
                    'status': 'success',
                    'format': 'pdf'
                }
            else:
                return {
                    'file_path': '',
                    'filename': filename,
                    'file_size': 0,
                    'status': 'error',
                    'error': 'PDF conversion failed',
                    'format': 'pdf'
                }
                
        except Exception as e:
            logger.error(f"Error generating advanced PDF: {e}")
            return {
                'file_path': '',
                'filename': filename or 'error.pdf',
                'file_size': 0,
                'status': 'error',
                'error': str(e),
                'format': 'pdf'
            }
    
    def generate_docx_advanced(self, resume_text: str, filename: str = None, style: str = 'modern') -> Dict[str, str]:
        """Generate DOCX using advanced HTML conversion"""
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"resume_{timestamp}.docx"
            
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Convert markdown to HTML
            html_content = self.markdown_to_html(resume_text)
            
            # Convert HTML to DOCX
            success = self.html_to_docx(html_content, file_path)
            
            if success and os.path.exists(file_path):
                return {
                    'file_path': file_path,
                    'filename': filename,
                    'file_size': os.path.getsize(file_path),
                    'status': 'success',
                    'format': 'docx'
                }
            else:
                return {
                    'file_path': '',
                    'filename': filename,
                    'file_size': 0,
                    'status': 'error',
                    'error': 'DOCX conversion failed',
                    'format': 'docx'
                }
                
        except Exception as e:
            logger.error(f"Error generating advanced DOCX: {e}")
            return {
                'file_path': '',
                'filename': filename or 'error.docx',
                'file_size': 0,
                'status': 'error',
                'error': str(e),
                'format': 'docx'
            }
    
    def generate_both_formats_advanced(self, resume_text: str, base_filename: str, style: str = 'modern') -> Dict[str, Dict]:
        """Generate both PDF and DOCX using advanced conversion"""
        pdf_result = self.generate_pdf_advanced(resume_text, f"{base_filename}.pdf", style)
        docx_result = self.generate_docx_advanced(resume_text, f"{base_filename}.docx", style)
        
        return {
            'pdf': pdf_result,
            'docx': docx_result,
            'base_filename': base_filename,
            'style': style,
            'generated_at': datetime.now().isoformat()
        }