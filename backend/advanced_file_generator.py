"""
Advanced File Generator with Proper Markdown Support
Uses markdown2 and weasyprint for high-quality document conversion
"""

import os
import io
import logging
import tempfile
from typing import Dict, Optional
from datetime import datetime
import markdown2
import weasyprint
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

class AdvancedFileGenerator:
    """Advanced file generator with proper markdown support"""
    
    def __init__(self, output_dir: str = "outputs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        # Configure markdown2 with extensions
        self.markdown_extras = [
            'fenced-code-blocks',
            'break-on-newline',
            'markdown-in-html',
            'tables',
            'strike',
            'task_list'
        ]
        
        # CSS for PDF styling
        self.pdf_css = """
        @page {
            size: A4;
            margin: 1in;
        }
        
        body {
            font-family: 'Arial', 'Helvetica', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            font-size: 12pt;
        }
        
        h1, h2, h3, h4, h5, h6 {
            font-weight: bold;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            color: #2c3e50;
        }
        
        h1 { font-size: 18pt; text-align: center; border-bottom: 2px solid #2c3e50; padding-bottom: 10px; }
        h2 { font-size: 16pt; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }
        h3 { font-size: 14pt; }
        h4 { font-size: 13pt; }
        
        p {
            margin-bottom: 1em;
            text-align: justify;
        }
        
        ul, ol {
            margin-bottom: 1em;
            padding-left: 1.5em;
        }
        
        li {
            margin-bottom: 0.3em;
        }
        
        strong, b {
            font-weight: bold;
            color: #2c3e50;
        }
        
        em, i {
            font-style: italic;
        }
        
        .contact-info {
            text-align: center;
            margin-bottom: 2em;
            border-bottom: 1px solid #bdc3c7;
            padding-bottom: 1em;
        }
        
        .section {
            margin-bottom: 1.5em;
        }
        
        .header-name {
            font-size: 24pt;
            font-weight: bold;
            color: #2c3e50;
        }
        
        .header-title {
            font-size: 14pt;
            color: #7f8c8d;
            margin-top: 0.2em;
        }
        """
    
    def clean_markdown_text(self, text: str) -> str:
        """Clean and prepare markdown text for conversion"""
        # Remove the prefix text that multi-agent adds
        if "Here is the final, document-ready content" in text:
            parts = text.split("Here is the final, document-ready content that will generate perfect files:")
            if len(parts) > 1:
                text = parts[1].strip()
        
        # Remove extra asterisks that cause issues
        text = re.sub(r'\*{3,}', '**', text)  # Replace *** with **
        
        # Fix header formatting - ensure proper spacing
        text = re.sub(r'^(#{1,6})\s*(.+)$', r'\1 \2', text, flags=re.MULTILINE)
        
        # Ensure bullet points have proper spacing
        text = re.sub(r'^\s*[\*\+\-]\s+', '* ', text, flags=re.MULTILINE)
        
        # Fix bold formatting around colons and special characters
        text = re.sub(r'\*\*([^*]+?)\*\*:', r'**\1:**', text)
        
        # Add proper line breaks before headers
        text = re.sub(r'\n(#{1,6}\s)', r'\n\n\1', text)
        
        return text.strip()
    
    def generate_pdf_from_markdown(self, markdown_text: str, filename: str = None) -> Dict[str, str]:
        """Generate PDF using weasyprint from markdown"""
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"resume_{timestamp}.pdf"
            
            if not filename.endswith('.pdf'):
                filename += '.pdf'
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Clean the markdown text
            clean_text = self.clean_markdown_text(markdown_text)
            
            # Convert markdown to HTML
            html_content = markdown2.markdown(clean_text, extras=self.markdown_extras)
            
            # Wrap in complete HTML document
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>Resume</title>
                <style>{self.pdf_css}</style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Generate PDF using weasyprint
            html_doc = weasyprint.HTML(string=full_html)
            html_doc.write_pdf(file_path)
            
            return {
                'file_path': file_path,
                'filename': filename,
                'file_size': os.path.getsize(file_path),
                'status': 'success',
                'format': 'pdf'
            }
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            return {
                'file_path': '',
                'filename': filename or 'error.pdf',
                'file_size': 0,
                'status': 'error',
                'error': str(e),
                'format': 'pdf'
            }
    
    def generate_docx_from_markdown(self, markdown_text: str, filename: str = None) -> Dict[str, str]:
        """Generate DOCX using python-docx with proper markdown parsing"""
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"resume_{timestamp}.docx"
            
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Clean the markdown text
            clean_text = self.clean_markdown_text(markdown_text)
            
            # Convert to HTML first for easier parsing
            html_content = markdown2.markdown(clean_text, extras=self.markdown_extras)
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Create new document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Process the HTML content
            self._add_content_from_html(doc, soup)
            
            # Save document
            doc.save(file_path)
            
            return {
                'file_path': file_path,
                'filename': filename,
                'file_size': os.path.getsize(file_path),
                'status': 'success',
                'format': 'docx'
            }
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            return {
                'file_path': '',
                'filename': filename or 'error.docx',
                'file_size': 0,
                'status': 'error',
                'error': str(e),
                'format': 'docx'
            }
    
    def _add_content_from_html(self, doc, soup):
        """Add content to DOCX document from HTML soup"""
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Add header
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text().strip(), level=min(level, 3))
                
                # Style the heading
                if level == 1:
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in heading.runs:
                        run.font.size = Pt(18)
                        run.font.bold = True
                elif level == 2:
                    for run in heading.runs:
                        run.font.size = Pt(14)
                        run.font.bold = True
                
            elif element.name == 'p':
                # Add paragraph
                if element.get_text().strip():
                    para = doc.add_paragraph()
                    self._add_formatted_text(para, element)
                    
            elif element.name in ['ul', 'ol']:
                # Add list items
                for li in element.find_all('li'):
                    if li.get_text().strip():
                        para = doc.add_paragraph(style='List Bullet')
                        self._add_formatted_text(para, li)
    
    def _add_formatted_text(self, paragraph, element):
        """Add formatted text to paragraph, handling bold, italic, etc."""
        for content in element.contents:
            if hasattr(content, 'name'):
                if content.name == 'strong' or content.name == 'b':
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name == 'em' or content.name == 'i':
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                else:
                    paragraph.add_run(content.get_text())
            else:
                # Plain text
                paragraph.add_run(str(content))
    
    def generate_both_formats(self, markdown_text: str, base_filename: str, style: str = 'modern') -> Dict:
        """Generate both PDF and DOCX from markdown"""
        results = {}
        
        # Generate PDF
        pdf_result = self.generate_pdf_from_markdown(markdown_text, f"{base_filename}.pdf")
        results['pdf'] = pdf_result
        
        # Generate DOCX
        docx_result = self.generate_docx_from_markdown(markdown_text, f"{base_filename}.docx")
        results['docx'] = docx_result
        
        # Add metadata
        results['base_filename'] = base_filename
        results['style'] = style
        results['generated_at'] = datetime.now().isoformat()
        
        return results