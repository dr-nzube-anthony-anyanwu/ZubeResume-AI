"""
File Generator Module
Converts tailored resume text to downloadable formats (PDF, DOCX)
Enhanced with Hugging Face text processing integration
"""

import os
import io
import logging
from typing import Dict, Optional, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfgen import canvas
import re

# Import text processor for final formatting
try:
    from .text_processor import TextProcessor
    from .markdown_converter import MarkdownConverter
except ImportError:
    from text_processor import TextProcessor
    from markdown_converter import MarkdownConverter

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileGenerator:
    """Class to generate downloadable resume files in various formats"""
    
    def __init__(self, output_dir: str = "outputs"):
        """
        Initialize the file generator
        
        Args:
            output_dir (str): Directory to save generated files
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        # Initialize text processor for final formatting
        try:
            self.text_processor = TextProcessor()
            self.markdown_converter = MarkdownConverter()
            logger.info("Text processor and markdown converter initialized in file generator")
        except Exception as e:
            logger.warning(f"Could not initialize text processor in file generator: {e}")
            self.text_processor = None
        
        # Document styling options
        self.styles = {
            'modern': {
                'font_family': 'Calibri',
                'header_color': colors.HexColor('#2C3E50'),
                'accent_color': colors.HexColor('#3498DB')
            },
            'classic': {
                'font_family': 'Times-Roman',
                'header_color': colors.black,
                'accent_color': colors.HexColor('#1F4E79')
            },
            'minimal': {
                'font_family': 'Helvetica',
                'header_color': colors.HexColor('#34495E'),
                'accent_color': colors.HexColor('#95A5A6')
            }
        }
    
    def _ensure_proper_formatting(self, resume_text: str) -> str:
        """Ensure text is properly formatted before file generation"""
        if not resume_text:
            return ""
        
        # Use markdown converter to clean up multi-agent output
        if hasattr(self, 'markdown_converter') and self.markdown_converter:
            try:
                logger.info("Converting markdown formatting to clean text...")
                clean_text = self.markdown_converter.convert_to_plain_formatted_text(resume_text)
                logger.info("Markdown conversion completed successfully")
                return clean_text
            except Exception as e:
                logger.warning(f"Markdown conversion failed: {e}")
        
        # Apply text processor for final formatting if available
        if self.text_processor:
            try:
                logger.info("Applying final text processing before file generation...")
                processed_text = self.text_processor.clean_and_structure_text(resume_text)
                logger.info("Final text processing completed successfully")
                return processed_text
            except Exception as e:
                logger.warning(f"Final text processing failed: {e}")
        
        # Fallback: basic formatting fixes if text processor not available
        return self._basic_formatting_fixes(resume_text)
    
    def _basic_formatting_fixes(self, text: str) -> str:
        """Basic formatting fixes as fallback"""
        # Fix common spacing issues
        text = re.sub(r'([a-zA-Z]),([a-zA-Z])', r'\1, \2', text)  # Comma spacing
        text = re.sub(r'([a-z])\.([A-Z])', r'\1. \2', text)  # Period spacing  
        text = re.sub(r'([a-zA-Z])\|([a-zA-Z])', r'\1 | \2', text)  # Pipe spacing
        text = re.sub(r'([a-zA-Z]):([a-zA-Z])', r'\1: \2', text)  # Colon spacing
        text = re.sub(r'(\d)([A-Za-z])', r'\1 \2', text)  # Number-letter spacing
        text = re.sub(r'(\d)%([A-Za-z])', r'\1% \2', text)  # Percentage spacing
        
        return text
    
    def generate_docx(self, resume_text: str, filename: str = None, 
                     style: str = 'modern') -> Dict[str, str]:
        """
        Generate a DOCX file from resume text
        
        Args:
            resume_text (str): The formatted resume text
            filename (str): Optional filename, auto-generated if not provided
            style (str): Style template to use
            
        Returns:
            Dict containing file path and generation status
        """
        try:
            # ENHANCED: Ensure proper formatting before processing
            resume_text = self._ensure_proper_formatting(resume_text)
            
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"tailored_resume_{timestamp}.docx"
            
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Create a new document
            doc = Document()
            
            # Set up document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Parse resume sections
            sections_data = self._parse_resume_sections(resume_text)
            
            # Add content to document
            self._add_docx_header(doc, sections_data.get('header', ''))
            self._add_docx_section(doc, 'SUMMARY', sections_data.get('summary', ''), is_summary=True)
            self._add_docx_section(doc, 'SKILLS', sections_data.get('skills', ''))
            self._add_docx_section(doc, 'EXPERIENCE', sections_data.get('experience', ''))
            self._add_docx_section(doc, 'EDUCATION', sections_data.get('education', ''))
            
            # Add any additional sections
            for section_name, content in sections_data.items():
                if section_name not in ['header', 'summary', 'skills', 'experience', 'education'] and content:
                    self._add_docx_section(doc, section_name.upper(), content)
            
            # Save the document
            doc.save(file_path)
            
            return {
                'file_path': file_path,
                'filename': filename,
                'file_size': os.path.getsize(file_path),
                'status': 'success',
                'format': 'docx'
            }
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            return {
                'file_path': '',
                'filename': filename or 'error.docx',
                'file_size': 0,
                'status': 'error',
                'error': str(e),
                'format': 'docx'
            }
    
    def generate_pdf(self, resume_text: str, filename: str = None, 
                    style: str = 'modern') -> Dict[str, str]:
        """
        Generate a PDF file from resume text
        
        Args:
            resume_text (str): The formatted resume text
            filename (str): Optional filename, auto-generated if not provided
            style (str): Style template to use
            
        Returns:
            Dict containing file path and generation status
        """
        try:
            # ENHANCED: Ensure proper formatting before processing
            resume_text = self._ensure_proper_formatting(resume_text)
            
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"tailored_resume_{timestamp}.pdf"
            
            if not filename.endswith('.pdf'):
                filename += '.pdf'
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(
                file_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=36,
                bottomMargin=36
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            style_config = self.styles.get(style, self.styles['modern'])
            
            # Create custom styles
            header_style = ParagraphStyle(
                'CustomHeader',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=style_config['header_color'],
                spaceAfter=12,
                alignment=1  # Center alignment
            )
            
            section_header_style = ParagraphStyle(
                'SectionHeader',
                parent=styles['Heading2'],
                fontSize=12,
                textColor=style_config['accent_color'],
                spaceBefore=12,
                spaceAfter=6,
                borderWidth=1,
                borderColor=style_config['accent_color'],
                borderPadding=3
            )
            
            # Parse resume sections
            sections_data = self._parse_resume_sections(resume_text)
            
            # Build story (content list)
            story = []
            
            # Add header with professional formatting
            if sections_data.get('header'):
                header_lines = sections_data['header'].split('\n')
                clean_lines = [line.strip() for line in header_lines if line.strip()]
                
                if clean_lines:
                    # Name (first line) - larger and bold
                    name_style = ParagraphStyle(
                        'NameStyle',
                        parent=styles['Title'],
                        fontSize=20,
                        spaceAfter=6,
                        alignment=1,  # Center alignment
                        textColor=colors.darkblue,
                        fontName='Helvetica-Bold'
                    )
                    story.append(Paragraph(clean_lines[0], name_style))
                    
                    # Contact info - smaller and centered
                    if len(clean_lines) > 1:
                        contact_style = ParagraphStyle(
                            'ContactStyle',
                            parent=styles['Normal'],
                            fontSize=11,
                            spaceAfter=12,
                            alignment=1,  # Center alignment
                            textColor=colors.darkgrey
                        )
                        contact_info = " | ".join(clean_lines[1:])
                        story.append(Paragraph(contact_info, contact_style))
                    
                    # Add separator line
                    separator_style = ParagraphStyle(
                        'SeparatorStyle',
                        parent=styles['Normal'],
                        fontSize=8,
                        spaceAfter=16,
                        alignment=1,
                        textColor=colors.lightgrey
                    )
                    story.append(Paragraph('_' * 60, separator_style))
            
            # Add sections
            section_order = ['summary', 'skills', 'experience', 'education']
            for section in section_order:
                content = sections_data.get(section, '')
                if content:
                    story.append(Paragraph(section.upper(), section_header_style))
                    self._add_pdf_section_content(story, content, styles)
                    story.append(Spacer(1, 6))
            
            # Add any additional sections
            for section_name, content in sections_data.items():
                if section_name not in ['header'] + section_order and content:
                    story.append(Paragraph(section_name.upper(), section_header_style))
                    self._add_pdf_section_content(story, content, styles)
                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            return {
                'file_path': file_path,
                'filename': filename,
                'file_size': os.path.getsize(file_path),
                'status': 'success',
                'format': 'pdf'
            }
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            return {
                'file_path': '',
                'filename': filename or 'error.pdf',
                'file_size': 0,
                'status': 'error',
                'error': str(e),
                'format': 'pdf'
            }
    
    def _parse_resume_sections(self, resume_text: str) -> Dict[str, str]:
        """Parse resume text into sections with enhanced detection and proper spacing"""
        sections = {}
        
        # Clean the text first - fix newline issues and spacing
        resume_text = resume_text.replace('\r\n', '\n').replace('\r', '\n')
        resume_text = re.sub(r'\n\s*\n', '\n\n', resume_text)  # Normalize multiple newlines
        
        # Enhanced section patterns - more comprehensive
        section_patterns = {
            'summary': r'(?:PROFESSIONAL SUMMARY|SUMMARY|OBJECTIVE|PROFILE|ABOUT|CAREER SUMMARY)\s*:?\s*$',
            'skills': r'(?:SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES|TECHNOLOGIES|CORE TECHNICAL SKILLS|TECHNICAL COMPETENCIES)\s*:?\s*$',
            'experience': r'(?:EXPERIENCE|WORK EXPERIENCE|PROFESSIONAL EXPERIENCE|EMPLOYMENT|PROJECT EXPERIENCE|WORK HISTORY)\s*:?\s*$',
            'education': r'(?:EDUCATION|ACADEMIC BACKGROUND|QUALIFICATIONS|ACADEMIC QUALIFICATIONS|EDUCATIONAL BACKGROUND)\s*:?\s*$',
            'projects': r'(?:PROJECTS|KEY PROJECTS|NOTABLE PROJECTS|PERSONAL PROJECTS)\s*:?\s*$',
            'certifications': r'(?:CERTIFICATIONS|CERTIFICATES|PROFESSIONAL CERTIFICATIONS|CREDENTIALS)\s*:?\s*$',
            'achievements': r'(?:ACHIEVEMENTS|ACCOMPLISHMENTS|AWARDS|HONORS|RESEARCH INTERESTS)\s*:?\s*$'
        }
        
        # Split into lines with better handling
        lines = []
        for line in resume_text.split('\n'):
            line = line.strip()
            if line:
                # Fix excessive spacing within lines
                line = re.sub(r'\s+', ' ', line)
                lines.append(line)
            else:
                # Preserve intentional blank lines for spacing
                lines.append('')
        
        # Parse sections
        current_section = 'header'
        current_content = []
        found_first_section = False
        
        for i, line in enumerate(lines):
            # Check for section headers
            is_section_header = False
            matched_section = None
            
            if line.strip():  # Only check non-empty lines
                line_upper = line.upper().strip()
                for section_name, pattern in section_patterns.items():
                    if re.match(pattern, line_upper):
                        is_section_header = True
                        matched_section = section_name
                        found_first_section = True
                        break
            
            if is_section_header:
                # Save previous section
                if current_section == 'header' and current_content:
                    sections['header'] = self._clean_header_content(current_content)
                elif current_content:
                    sections[current_section] = self._clean_section_content(current_content, current_section)
                
                # Start new section
                current_section = matched_section
                current_content = []
                
            elif not found_first_section:
                # Still in header section
                current_content.append(line)
            else:
                # Content for current section
                current_content.append(line)
        
        # Save the last section
        if current_section == 'header' and current_content:
            sections['header'] = self._clean_header_content(current_content)
        elif current_content:
            sections[current_section] = self._clean_section_content(current_content, current_section)
        
        return sections
    
    def _clean_header_content(self, content_lines: list) -> str:
        """Clean and format header content"""
        if not content_lines:
            return ""
        
        # Separate name from contact info
        cleaned_lines = []
        for line in content_lines:
            line = line.strip()
            if line:
                # Remove titles like "Dr." if they're standalone
                if line in ['Dr.', 'MD', 'PhD', 'O.D']:
                    continue
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _clean_section_content(self, content_lines: list, section_type: str) -> str:
        """Clean and format section content with proper spacing and line breaks"""
        if not content_lines:
            return ""
        
        cleaned_lines = []
        i = 0
        
        while i < len(content_lines):
            line = content_lines[i].strip()
            
            # Skip empty lines in input but preserve spacing structure
            if not line:
                # Add controlled spacing
                if cleaned_lines and cleaned_lines[-1] != "":
                    cleaned_lines.append("")
                i += 1
                continue
            
            # Handle different section types with enhanced formatting
            if section_type == 'skills':
                # Handle skill categories with proper line breaks
                if ':' in line and not line.startswith(('•', '-', '*', '●')):
                    # This is a category line (e.g., "Programming & Tools: Python, Java")
                    if cleaned_lines and cleaned_lines[-1] != "":
                        cleaned_lines.append("")  # Add space before category
                    cleaned_lines.append(line)
                    # Look ahead to see if next line is another category
                    if i + 1 < len(content_lines):
                        next_line = content_lines[i + 1].strip()
                        if next_line and ':' in next_line and not next_line.startswith(('•', '-', '*', '●')):
                            cleaned_lines.append("")  # Add space between categories
                else:
                    cleaned_lines.append(line)
                    
            elif section_type == 'experience':
                # Enhanced experience section formatting
                if line.startswith('●') or '—' in line:
                    # Project title with description - add space before
                    if cleaned_lines and cleaned_lines[-1] != "":
                        cleaned_lines.append("")
                    cleaned_lines.append(line)
                    
                elif ('engineer' in line.lower() or 'developer' in line.lower() or 
                      'analyst' in line.lower() or 'manager' in line.lower() or 
                      'lead' in line.lower() or 'specialist' in line.lower()):
                    # Role/position line
                    cleaned_lines.append(line)
                    # Add space after role line for readability
                    cleaned_lines.append("")
                    
                elif line.startswith(('•', '-', '*')):
                    # Achievement bullet point
                    cleaned_lines.append(line)
                    
                else:
                    # Company name, dates, or other descriptive text
                    cleaned_lines.append(line)
                    
            elif section_type == 'summary':
                # Summary should flow as paragraphs with good spacing
                if len(line) > 10:  # Avoid single words on their own line
                    cleaned_lines.append(line)
                    # Add paragraph breaks for readability
                    if i + 1 < len(content_lines) and len(content_lines[i + 1].strip()) > 10:
                        if not line.endswith('.'):
                            cleaned_lines.append("")  # Add space between summary points
                            
            elif section_type in ['achievements', 'certifications', 'projects']:
                # Ensure bullet points are properly formatted
                if line and not line.startswith(('•', '-', '*', '●')):
                    # Add bullet if it's a list item
                    if section_type in ['achievements', 'certifications']:
                        line = f"• {line}"
                
                # Add spacing between major items
                if line.startswith('●') or (section_type == 'projects' and len(line) > 30):
                    if cleaned_lines and cleaned_lines[-1] != "":
                        cleaned_lines.append("")
                        
                cleaned_lines.append(line)
                
            elif section_type == 'education':
                # Education formatting with proper spacing
                if any(word in line.lower() for word in ['university', 'college', 'institute', 'school', 'bachelor', 'master', 'phd']):
                    # Major education entry - add space before
                    if cleaned_lines and cleaned_lines[-1] != "":
                        cleaned_lines.append("")
                        
                cleaned_lines.append(line)
                
            else:
                # Default formatting - preserve structure but clean spacing
                cleaned_lines.append(line)
            
            i += 1
        
        # Final cleanup - remove excessive blank lines but preserve structure
        final_lines = []
        blank_line_count = 0
        
        for line in cleaned_lines:
            if line == "":
                blank_line_count += 1
                if blank_line_count <= 1:  # Allow max 1 consecutive blank line
                    final_lines.append(line)
            else:
                blank_line_count = 0
                final_lines.append(line)
        
        # Remove trailing blank lines
        while final_lines and final_lines[-1] == "":
            final_lines.pop()
            
        return '\n'.join(final_lines)
    
    def _add_docx_header(self, doc: Document, header_text: str):
        """Add header section to DOCX document with professional formatting and proper spacing"""
        if not header_text:
            return
        
        lines = [line.strip() for line in header_text.split('\n') if line.strip()]
        
        if lines:
            # First line is usually the name - make it prominent
            name_paragraph = doc.add_paragraph()
            name_run = name_paragraph.add_run(lines[0])
            name_run.font.size = Pt(22)
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(0x1F, 0x4F, 0x7F)  # Professional dark blue
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_paragraph.space_after = Pt(6)
            
            # Extract and format contact information
            contact_info = []
            for line in lines[1:]:
                if line.strip():
                    contact_info.append(line.strip())
            
            if contact_info:
                # Format contact info with icons/separators
                contact_paragraph = doc.add_paragraph()
                contact_text = " | ".join(contact_info)
                contact_run = contact_paragraph.add_run(contact_text)
                contact_run.font.size = Pt(11)
                contact_run.font.color.rgb = RGBColor(0x4F, 0x4F, 0x4F)  # Gray
                contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_paragraph.space_after = Pt(16)
            
            # Add a subtle separator line
            separator_paragraph = doc.add_paragraph()
            separator_run = separator_paragraph.add_run('_' * 80)
            separator_run.font.size = Pt(8)
            separator_run.font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)  # Light gray
            separator_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            separator_paragraph.space_after = Pt(16)
    
    def _add_docx_section(self, doc: Document, title: str, content: str, is_summary: bool = False):
        """Add a section to DOCX document with enhanced spacing and formatting"""
        if not content:
            return
        
        # Add section title with proper spacing
        section_title = doc.add_paragraph()
        section_title.space_before = Pt(16)
        section_title.space_after = Pt(8)
        title_run = section_title.add_run(title)
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1F, 0x4F, 0x7F)  # Professional blue for headers
        
        # Process content with enhanced formatting
        if is_summary:
            # Summary as flowing paragraphs
            paragraphs = content.split('\n\n')  # Split on double newlines
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # Clean and join lines within a paragraph
                    lines = [line.strip() for line in paragraph_text.split('\n') if line.strip()]
                    combined_text = ' '.join(lines)
                    
                    summary_paragraph = doc.add_paragraph(combined_text)
                    summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    summary_paragraph.space_after = Pt(12)
                    summary_paragraph.paragraph_format.line_spacing = 1.15
                    
                    # Ensure summary text is black
                    for run in summary_paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        else:
            # Process other sections line by line with proper spacing
            lines = content.split('\n')
            
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                # Handle empty lines for spacing
                if not line:
                    # Add controlled spacing
                    spacer = doc.add_paragraph()
                    spacer.space_after = Pt(6)
                    i += 1
                    continue
                
                # Skip redundant section headers
                if line.upper() in [title.upper(), title.upper() + ':']: 
                    i += 1
                    continue
                
                # Handle bullet points with proper indentation
                if line.startswith(('•', '-', '*')):
                    clean_line = line.lstrip('•-* ').strip()
                    if clean_line:
                        bullet_paragraph = doc.add_paragraph()
                        bullet_paragraph.style = 'List Bullet'
                        bullet_run = bullet_paragraph.add_run(clean_line)
                        bullet_run.font.size = Pt(11)
                        bullet_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black text
                        bullet_paragraph.space_after = Pt(4)
                        bullet_paragraph.paragraph_format.left_indent = Inches(0.25)
                        bullet_paragraph.paragraph_format.line_spacing = 1.1
                
                # Handle project titles (● Title — Description format)
                elif line.startswith('●') or '—' in line:
                    project_paragraph = doc.add_paragraph()
                    project_paragraph.space_before = Pt(8)
                    project_paragraph.space_after = Pt(4)
                    
                    if '—' in line:
                        # Split project name and description
                        parts = line.split('—', 1)
                        project_name = parts[0].strip().lstrip('●').strip()
                        project_desc = parts[1].strip()
                        
                        # Project name in bold blue
                        name_run = project_paragraph.add_run(f"● {project_name}")
                        name_run.font.bold = True
                        name_run.font.size = Pt(12)
                        name_run.font.color.rgb = RGBColor(0x1F, 0x4F, 0x7F)
                        
                        # Description in normal black text
                        if project_desc:
                            desc_run = project_paragraph.add_run(f" — {project_desc}")
                            desc_run.font.size = Pt(11)
                            desc_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    else:
                        # Simple project title
                        project_run = project_paragraph.add_run(line)
                        project_run.font.bold = True
                        project_run.font.size = Pt(12)
                        project_run.font.color.rgb = RGBColor(0x1F, 0x4F, 0x7F)
                
                # Handle role/position lines with tech stack
                elif ('|' in line and 'tech:' in line.lower()) or any(role in line.lower() for role in ['engineer', 'developer', 'analyst', 'manager', 'lead', 'specialist']):
                    role_paragraph = doc.add_paragraph()
                    role_paragraph.space_after = Pt(6)
                    
                    if '|' in line and 'tech:' in line.lower():
                        # Split role and tech stack
                        parts = line.split('|', 1)
                        role_part = parts[0].strip()
                        tech_part = parts[1].strip()
                        
                        # Role in bold black
                        role_run = role_paragraph.add_run(role_part)
                        role_run.font.bold = True
                        role_run.font.size = Pt(11)
                        role_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                        
                        # Tech stack in italic gray  
                        if tech_part:
                            tech_run = role_paragraph.add_run(f" | {tech_part}")
                            tech_run.font.italic = True
                            tech_run.font.size = Pt(10)
                            tech_run.font.color.rgb = RGBColor(0x4F, 0x4F, 0x4F)
                    else:
                        # Simple role line
                        role_run = role_paragraph.add_run(line)
                        role_run.font.bold = True
                        role_run.font.size = Pt(11)
                        role_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                
                # Handle skills categories (Programming & Tools: Python, Java, etc.)
                elif ':' in line and any(skill_word in line.lower() for skill_word in ['programming', 'machine', 'cloud', 'data', 'tools', 'platforms', 'ethics', 'languages']):
                    skill_paragraph = doc.add_paragraph()
                    skill_paragraph.space_after = Pt(6)
                    
                    # Split at colon
                    parts = line.split(':', 1)
                    category = parts[0].strip()
                    skills = parts[1].strip() if len(parts) > 1 else ''
                    
                    # Category in bold black
                    cat_run = skill_paragraph.add_run(f"{category}:")
                    cat_run.font.bold = True
                    cat_run.font.size = Pt(11)
                    cat_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    
                    # Skills in normal black text
                    if skills:
                        skill_run = skill_paragraph.add_run(f" {skills}")
                        skill_run.font.size = Pt(11)
                        skill_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                
                # Handle education entries
                elif any(edu_word in line.lower() for edu_word in ['university', 'college', 'bachelor', 'master', 'phd', 'degree']):
                    edu_paragraph = doc.add_paragraph()
                    edu_paragraph.space_before = Pt(6)
                    edu_paragraph.space_after = Pt(4)
                    
                    # Education institution/degree in bold
                    edu_run = edu_paragraph.add_run(line)
                    edu_run.font.bold = True
                    edu_run.font.size = Pt(11)
                    edu_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                
                # Regular content lines
                else:
                    regular_paragraph = doc.add_paragraph(line)
                    regular_paragraph.space_after = Pt(6)
                    regular_paragraph.paragraph_format.line_spacing = 1.1
                    
                    # Ensure all text is black and properly sized
                    for run in regular_paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                
                i += 1
        
        # Add proper spacing after section
        spacer_paragraph = doc.add_paragraph()
        spacer_paragraph.space_after = Pt(16)
    
    def _add_pdf_section_content(self, story: list, content: str, styles):
        """Add section content to PDF story with enhanced formatting and proper spacing"""
        # Create enhanced styles for better formatting
        bullet_style = ParagraphStyle(
            'BulletStyle',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=24,
            bulletIndent=12,
            spaceAfter=4,
            spaceBefore=2,
            leading=14
        )
        
        normal_style = ParagraphStyle(
            'NormalStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leading=14
        )
        
        project_title_style = ParagraphStyle(
            'ProjectTitleStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceBefore=8,
            spaceAfter=4,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold',
            leading=16
        )
        
        role_style = ParagraphStyle(
            'RoleStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            fontName='Helvetica-Bold',
            leading=13
        )
        
        tech_style = ParagraphStyle(
            'TechStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            fontName='Helvetica-Oblique',
            textColor=colors.darkgrey,
            leading=13
        )
        
        skills_category_style = ParagraphStyle(
            'SkillsCategoryStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            fontName='Helvetica-Bold',
            leading=14
        )
        
        # Process content line by line with proper spacing
        lines = content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Handle empty lines for spacing
            if not line:
                story.append(Spacer(1, 6))  # Add controlled spacing
                i += 1
                continue
            
            # Handle bullet points with proper formatting
            if line.startswith(('•', '-', '*')):
                clean_line = line.lstrip('•-* ').strip()
                if clean_line:
                    story.append(Paragraph(f"• {clean_line}", bullet_style))
            
            # Handle project titles (lines starting with ● or containing —)
            elif line.startswith('●') or '—' in line:
                if '—' in line:
                    parts = line.split('—', 1)
                    project_name = parts[0].strip().lstrip('●').strip()
                    project_desc = parts[1].strip()
                    formatted_line = f"<b>● {project_name}</b> — {project_desc}"
                else:
                    formatted_line = f"<b>{line}</b>"
                story.append(Paragraph(formatted_line, project_title_style))
            
            # Handle role/position lines with tech stack
            elif ('|' in line and 'tech:' in line.lower()) or any(role in line.lower() for role in ['engineer', 'developer', 'analyst', 'manager', 'lead', 'specialist']):
                if '|' in line and 'tech:' in line.lower():
                    parts = line.split('|', 1)
                    role_part = parts[0].strip()
                    tech_part = parts[1].strip()
                    
                    story.append(Paragraph(f"<b>{role_part}</b>", role_style))
                    if tech_part:
                        story.append(Paragraph(f"<i>{tech_part}</i>", tech_style))
                else:
                    story.append(Paragraph(f"<b>{line}</b>", role_style))
            
            # Handle skills categories (Programming & Tools:, Machine Learning:, etc.)
            elif ':' in line and any(skill_word in line.lower() for skill_word in ['programming', 'machine', 'cloud', 'data', 'tools', 'platforms', 'ethics', 'languages']):
                parts = line.split(':', 1)
                category = parts[0].strip()
                skills = parts[1].strip() if len(parts) > 1 else ''
                
                if skills:
                    formatted_line = f"<b>{category}:</b> {skills}"
                    story.append(Paragraph(formatted_line, skills_category_style))
                else:
                    story.append(Paragraph(f"<b>{category}:</b>", skills_category_style))
            
            # Handle education entries
            elif any(edu_word in line.lower() for edu_word in ['university', 'college', 'bachelor', 'master', 'phd', 'degree']):
                story.append(Paragraph(f"<b>{line}</b>", role_style))
            
            # Regular content lines
            else:
                story.append(Paragraph(line, normal_style))
            
            i += 1
        
        # Add spacing after section
        story.append(Spacer(1, 12))
    
    def generate_both_formats(self, resume_text: str, base_filename: str = None, 
                            style: str = 'modern') -> Dict:
        """
        Generate both PDF and DOCX formats
        
        Args:
            resume_text (str): The formatted resume text
            base_filename (str): Base filename (without extension)
            style (str): Style template to use
            
        Returns:
            Dict containing results for both formats
        """
        if not base_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"tailored_resume_{timestamp}"
        
        # Generate both formats
        docx_result = self.generate_docx(resume_text, f"{base_filename}.docx", style)
        pdf_result = self.generate_pdf(resume_text, f"{base_filename}.pdf", style)
        
        return {
            'docx': docx_result,
            'pdf': pdf_result,
            'base_filename': base_filename,
            'style': style,
            'generated_at': datetime.now().isoformat()
        }
    
    def generate_both_formats_advanced(self, resume_text: str, base_filename: str = None, 
                            style: str = 'modern') -> Dict:
        """
        Advanced generate both formats - alias for compatibility with V2 generator
        """
        return self.generate_both_formats(resume_text, base_filename, style)
    
    def get_file_bytes(self, file_path: str) -> bytes:
        """Get file content as bytes for download"""
        try:
            with open(file_path, 'rb') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {str(e)}")
            raise
    
    def cleanup_old_files(self, max_age_hours: int = 24):
        """Clean up old generated files"""
        try:
            current_time = datetime.now().timestamp()
            
            for filename in os.listdir(self.output_dir):
                file_path = os.path.join(self.output_dir, filename)
                if os.path.isfile(file_path):
                    file_age = current_time - os.path.getmtime(file_path)
                    if file_age > (max_age_hours * 3600):  # Convert hours to seconds
                        os.remove(file_path)
                        logger.info(f"Removed old file: {filename}")
        
        except Exception as e:
            logger.error(f"Error cleaning up files: {str(e)}")

def test_file_generator():
    """Test function for the file generator"""
    generator = FileGenerator()
    
    sample_resume = """
John Doe
john.doe@email.com | (555) 123-4567 | LinkedIn: linkedin.com/in/johndoe

SUMMARY
Experienced software developer with 5 years of experience in full-stack development.
Skilled in Python, JavaScript, and cloud technologies.

SKILLS
• Programming Languages: Python, JavaScript, Java, SQL
• Frameworks: React, Django, Flask, Node.js
• Cloud Platforms: AWS, Azure
• Databases: PostgreSQL, MongoDB

EXPERIENCE
Senior Software Developer
Tech Corp Inc. | 2022 - Present
• Led development of web applications serving 100K+ users
• Implemented microservices architecture reducing response time by 40%
• Mentored junior developers and conducted code reviews

Software Developer
StartupXYZ | 2020 - 2022
• Developed and maintained e-commerce platform
• Collaborated with cross-functional teams
• Implemented automated testing increasing code coverage to 90%

EDUCATION
Bachelor of Science in Computer Science
University of Technology | 2016 - 2020
• Graduated Magna Cum Laude
• Relevant Coursework: Data Structures, Algorithms, Software Engineering
"""
    
    # Test both format generation
    result = generator.generate_both_formats(sample_resume, "test_resume")
    
    print("File Generation Test Results:")
    print(f"DOCX Status: {result['docx']['status']}")
    print(f"PDF Status: {result['pdf']['status']}")
    
    if result['docx']['status'] == 'success':
        print(f"DOCX File: {result['docx']['filename']} ({result['docx']['file_size']} bytes)")
    
    if result['pdf']['status'] == 'success':
        print(f"PDF File: {result['pdf']['filename']} ({result['pdf']['file_size']} bytes)")

if __name__ == "__main__":
    test_file_generator()